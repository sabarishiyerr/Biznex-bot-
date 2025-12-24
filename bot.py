import os
import datetime
import requests

from dotenv import load_dotenv

import gspread
from google.oauth2.service_account import Credentials

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

BOT_TIMEZONE_OFFSET_HOURS = 5
BOT_TIMEZONE_OFFSET_MINUTES = 30


# =========================
# 1. LOAD CONFIG / CLIENTS
# =========================

load_dotenv()

RUN_MODE = (os.getenv("RUN_MODE") or "simulate").lower()


GOOGLE_SHEETS_CRED_PATH = os.getenv("GOOGLE_SHEETS_CRED_PATH")
GOOGLE_SHEETS_DOC_NAME = os.getenv("GOOGLE_SHEETS_DOC_NAME")

FB_PAGE_ACCESS_TOKEN = os.getenv("FB_PAGE_ACCESS_TOKEN")
FB_PAGE_ID = os.getenv("FB_PAGE_ID")

# placeholders for future real integrations (currently unused / simulated)
LINKEDIN_ACCESS_TOKEN = os.getenv("LINKEDIN_ACCESS_TOKEN")
INSTAGRAM_ACCESS_TOKEN = os.getenv("INSTAGRAM_ACCESS_TOKEN")


def get_gspread_client():
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds = Credentials.from_service_account_file(
        GOOGLE_SHEETS_CRED_PATH, scopes=scopes
    )
    gc = gspread.authorize(creds)
    return gc


# =========================
# 2. SHEETS HELPERS
# =========================

def get_sheets():
    gc = get_gspread_client()
    sh = gc.open(GOOGLE_SHEETS_DOC_NAME)
    content_sheet = sh.worksheet("ContentPlan")
    log_sheet = sh.worksheet("PostLog")
    return content_sheet, log_sheet

def add_content_item(date, time, platforms, idea,
                     caption="", image_url="", hashtags="", groups=""):
    """
    Append a new content row to the ContentPlan sheet with status='pending'.

    date: string, ideally 'YYYY-MM-DD' (can be blank for 'any day')
    platforms: string like 'FB, IG, LinkedIn'
    idea: main idea / prompt for the post
    caption, image_url, hashtags, groups: optional strings
    """
    content_sheet, _ = get_sheets()
    records = content_sheet.get_all_records()

    # Calculate next ID (max existing id + 1)
    next_id = 1
    if records:
        existing_ids = [
            int(r["id"])
            for r in records
            if "id" in r and str(r["id"]).strip().isdigit()
        ]
        if existing_ids:
            next_id = max(existing_ids) + 1

    # Our header layout: id, date, platforms, idea, caption,
    # image_url, hashtags, groups, status
    new_row = [
        next_id,
        date,
        time,
        platforms,
        idea,
        caption,
        image_url,
        hashtags,
        groups,
        "pending",
    ]

    content_sheet.append_row(new_row)
    return next_id

def normalize_sheet_date(value: str) -> str:
    """
    Convert different sheet date formats to YYYY-MM-DD.
    Returns "" if blank/unparseable.
    """
    if value is None:
        return ""
    s = str(value).strip()
    if not s:
        return ""

    # Already ISO date
    try:
        return datetime.date.fromisoformat(s).isoformat()
    except ValueError:
        pass

    # Common formats from Sheets: DD-MM-YYYY or DD/MM/YYYY
    for fmt in ("%d-%m-%Y", "%d/%m/%Y", "%Y/%m/%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue

    return ""

def normalize_sheet_time(value: str) -> str:
    """
    Normalize time values from the sheet into HH:MM (24-hour).
    Returns "" if blank/unparseable.
    """
    if value is None:
        return ""
    s = str(value).strip()
    if not s:
        return ""

    # Common time formats from Sheets
    for fmt in ("%H:%M", "%H:%M:%S"):
        try:
            return datetime.datetime.strptime(s, fmt).time().strftime("%H:%M")
        except ValueError:
            continue

    return ""


def find_all_pending_content(content_sheet):
    """
    Find ALL rows where status == 'pending'
    and scheduled date is today OR earlier (past).
    If date == today and a time is provided, only post when time <= now.
    Returns a list of row dicts with '__row_index__' added.
    """
    records = content_sheet.get_all_records()
    pending_rows = []

    utc_now = datetime.datetime.utcnow()
    now = utc_now + datetime.timedelta(
        hours=BOT_TIMEZONE_OFFSET_HOURS,
        minutes=BOT_TIMEZONE_OFFSET_MINUTES
    )

    today = now.date()
    current_time = now.time()

    for idx, row in enumerate(records, start=2):
        status = (row.get("status") or "").strip().lower()
        if status != "pending":
            continue

        # Normalize date/time from the sheet
        date_val = normalize_sheet_date(row.get("date"))   # returns "YYYY-MM-DD" or ""
        time_val = normalize_sheet_time(row.get("time"))   # returns "HH:MM" or ""

        # Optional debug: keep normalized values in row dict
        row["date"] = date_val
        row["time"] = time_val

        # Decide if this row is due
        if date_val:
            post_date = datetime.date.fromisoformat(date_val)

            # Future date -> not due yet
            if post_date > today:
                continue

            # Same date + time provided -> only due if time has passed
            if post_date == today and time_val:
                post_time = datetime.datetime.strptime(time_val, "%H:%M").time()
                if post_time > current_time:
                    continue

            # If post_date < today OR post_date == today (and time ok), it's due
        else:
            # Blank date means "post any day the bot runs"
            pass

        # If we reach here, row is due
        row["__row_index__"] = idx
        pending_rows.append(row)

    return pending_rows

def get_column_index_by_header(ws, header_name: str) -> int:
    """
    Returns the 1-based column index for a header name (case-insensitive).
    """
    headers = ws.row_values(1)
    headers_lower = [h.strip().lower() for h in headers]
    target = header_name.strip().lower()

    if target not in headers_lower:
        raise ValueError(f"Header '{header_name}' not found. Headers: {headers}")

    return headers_lower.index(target) + 1  # 1-based for gspread


def update_content_status(content_sheet, row_index, new_status):
    """
    Update the status cell for a given row using the 'status' header.
    """
    status_col = get_column_index_by_header(content_sheet, "status")
    content_sheet.update_cell(row_index, status_col, new_status)




def append_post_log(log_sheet, content_id, platform, caption_used, post_url):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_sheet.append_row([timestamp, content_id, platform, caption_used, post_url])


# =========================
# 3. CAPTION GENERATION (NO OPENAI)
# =========================

def generate_caption_if_needed(platform, idea, caption_existing, hashtags_existing):
    """
    TEMP VERSION (no OpenAI):
    - Uses caption/hashtags from the sheet if provided.
    - If they are empty, builds a simple fallback caption + basic hashtags.
    This lets us test the full pipeline without calling the OpenAI API.
    """
    caption = (caption_existing or "").strip()
    hashtags = (hashtags_existing or "").strip()

    # If no caption provided in sheet, create a simple one from the idea
    if not caption:
        caption = f"{idea} (auto-generated caption for {platform})"

    # If no hashtags provided, add some generic ones
    if not hashtags:
        hashtags = "#globalbiznex #marketing #automation"

    full_caption = caption
    if hashtags:
        full_caption += "\n\n" + hashtags

    return full_caption


# =========================
# 4. SOCIAL POSTING (SIMULATED FB / LINKEDIN / IG)
# =========================

""" def post_to_facebook(caption):
    
    Post text to a Facebook Page using the Graph API.

    If FB creds not set, just print and return a fake URL.
    
    if not FB_PAGE_ACCESS_TOKEN or not FB_PAGE_ID:
        print("[INFO] FB credentials not set. Simulating FB page post...")
        print("----- FB PAGE POST START -----")
        print(caption)
        print("----- FB PAGE POST END -------")
        fake_url = "https://facebook.com/fake_page_post"
        return fake_url

    # Real FB API call would go here
    url = f"https://graph.facebook.com/{FB_PAGE_ID}/feed"
    params = {
        "message": caption,
        "access_token": FB_PAGE_ACCESS_TOKEN,
    }

    resp = requests.post(url, data=params)
    if resp.status_code != 200:
        print("[ERROR] FB post failed:", resp.text)
        return None

    data = resp.json()
    post_id = data.get("id")
    if not post_id:
        return None

    post_url = f"https://www.facebook.com/{FB_PAGE_ID}/posts/{post_id.split('_')[-1]}"
    return post_url """

def post_to_facebook(caption):
    if RUN_MODE != "live":
        print("[SIMULATE] FB page post")
        return "https://facebook.com/fake_page_post"

    # live posting code (Graph API) below...



""" def post_to_linkedin(caption):
    Simulated LinkedIn post.
    Later, this can be replaced with a real LinkedIn API call.

    if not LINKEDIN_ACCESS_TOKEN:
        print("[INFO] LinkedIn credentials not set. Simulating LinkedIn post...")
        print("----- LINKEDIN POST START -----")
        print(caption)
        print("----- LINKEDIN POST END -------")
        fake_url = "https://linkedin.com/posts/fake_linkedin_post"
        return fake_url

    # Real LinkedIn API integration would go here.
    # For now, we just simulate even if token is set.
    fake_url = "https://linkedin.com/posts/fake_linkedin_post_real"
    return fake_url """

def post_to_linkedin(caption):
    if RUN_MODE != "live":
        print("[SIMULATE] LinkedIn post")
        return "https://linkedin.com/posts/fake_linkedin_post"



""" def post_to_instagram(caption):
    Simulated Instagram post.
    Later, this can be replaced with a real Instagram Graph API call.

    if not INSTAGRAM_ACCESS_TOKEN:
        print("[INFO] Instagram credentials not set. Simulating IG post...")
        print("----- INSTAGRAM POST START -----")
        print(caption)
        print("----- INSTAGRAM POST END -------")
        fake_url = "https://instagram.com/p/fake_instagram_post"
        return fake_url

    # Real Instagram API integration would go here.
    fake_url = "https://instagram.com/p/fake_instagram_post_real"
    return fake_url """

def post_to_instagram(caption):
    if RUN_MODE != "live":
        print("[SIMULATE] IG post")
        return "https://instagram.com/p/fake_instagram_post"



# =========================
# 5. WORD DOC LOGGING
# =========================

DOCX_LOG_PATH = "post_log.docx"


def log_to_word_doc(content_id, platform, caption, post_url):
    try:
        if os.path.exists(DOCX_LOG_PATH):
            # Try to open existing docx
            doc = Document(DOCX_LOG_PATH)
        else:
            # Create new docx
            doc = Document()
            doc.add_heading("Social Media Post Log", level=1)
    except PackageNotFoundError:
        # File exists but is not a valid .docx (corrupted or wrong type)
        print("[WARN] post_log.docx is not a valid Word file. Recreating it.")
        doc = Document()
        doc.add_heading("Social Media Post Log", level=1)

    # Add entry
    doc.add_paragraph(
        f"Time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    doc.add_paragraph(f"Content ID: {content_id}")
    doc.add_paragraph(f"Platform: {platform}")
    doc.add_paragraph("Caption:")
    doc.add_paragraph(caption)
    doc.add_paragraph(f"Post URL: {post_url}")
    doc.add_paragraph("-" * 40)

    # Try to save main log file
    try:
        doc.save(DOCX_LOG_PATH)
    except PermissionError:
        # File is probably open in Word, so save to a backup file instead
        backup_name = f"post_log_backup_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        print(f"[WARN] Could not write to '{DOCX_LOG_PATH}' (file in use).")
        print(f"[WARN] Saving this log entry to '{backup_name}' instead.")
        doc.save(backup_name)



# =========================
# 6. MAIN BOT LOGIC (MULTIPLE ROWS)
# =========================

def process_all_pending_items():
    content_sheet, log_sheet = get_sheets()

    pending_rows = find_all_pending_content(content_sheet)
    if not pending_rows:
        print("No pending content for today. Nothing to do.")
        return

    print(f"Found {len(pending_rows)} pending item(s).")

    for row in pending_rows:
        print("\n====================================")
        print("Processing row:", row)
        print("====================================")

        row_index = row["__row_index__"]
        content_id = row.get("id")
        idea = row.get("idea", "")
        caption_existing = row.get("caption", "")
        hashtags_existing = row.get("hashtags", "")
        platforms_raw = row.get("platforms", "")

        # split platforms like "FB, IG, LinkedIn"
        platforms = [p.strip() for p in platforms_raw.split(",") if p.strip()]

        if not platforms:
            print("No platforms specified; marking as 'no_platforms'.")
            update_content_status(content_sheet, row_index, "no_platforms")
            continue

        all_success = True

        for platform in platforms:
            full_caption = generate_caption_if_needed(
                platform, idea, caption_existing, hashtags_existing
            )

            print(f"\n--- Final caption for {platform} ---")
            print(full_caption)
            print("------------------------------------")

            post_url = None
            platform_lower = platform.lower()

            if platform_lower in ["fb", "facebook"]:
                post_url = post_to_facebook(full_caption)
            elif platform_lower in ["li", "linkedin"]:
                post_url = post_to_linkedin(full_caption)
            elif platform_lower in ["ig", "instagram"]:
                post_url = post_to_instagram(full_caption)
            else:
                print(f"[WARN] Platform '{platform}' not implemented yet. Skipping.")
                all_success = False
                continue

            if post_url:
                print(f"[OK] Posted to {platform}: {post_url}")
                # Log main post in Sheets and Word
                append_post_log(log_sheet, content_id, platform, full_caption, post_url)
                log_to_word_doc(content_id, platform, full_caption, post_url)

                # If this is Facebook, also handle group shares (simulated)
                if platform_lower in ["fb", "facebook"]:
                    groups_raw = (row.get("groups") or "").strip()
                    if groups_raw:
                        groups = [g.strip() for g in groups_raw.split(",") if g.strip()]
                        for group_name in groups:
                            fake_group_url = (
                                f"https://facebook.com/groups/"
                                f"{group_name.replace(' ', '_')}/fake_post"
                            )
                            print(
                                f"[INFO] Simulating share to group '{group_name}': "
                                f"{fake_group_url}"
                            )

                            append_post_log(
                                log_sheet,
                                content_id,
                                f"FB-Group: {group_name}",
                                full_caption,
                                fake_group_url,
                            )
                            log_to_word_doc(
                                content_id,
                                f"FB-Group: {group_name}",
                                full_caption,
                                fake_group_url,
                            )
                    else:
                        print("[INFO] No groups specified for this row; skipping group shares.")
            else:
                print(f"[ERROR] Failed to get post URL for {platform}")
                all_success = False

        new_status = "posted" if all_success else "partial"
        update_content_status(content_sheet, row_index, new_status)
        print(f"Updated content ID {content_id} to status '{new_status}'.")


if __name__ == "__main__":
    process_all_pending_items()
